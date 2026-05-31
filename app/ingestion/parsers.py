"""
Document parsers — one per file type, unified interface.
Each parser returns clean normalized text ready for chunking.
"""
import re
import unicodedata
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Optional

import aiofiles
from app.core.exceptions import IngestionError, UnsupportedFileTypeError
from app.core.logging import get_logger
from app.core.models import FileType

logger = get_logger(__name__)


# ─── Text Normalizer ──────────────────────────────────────────────────────────

class TextNormalizer:
    """Cleans and normalizes raw extracted text."""

    _MULTI_NEWLINE = re.compile(r"\n{3,}")
    _MULTI_SPACE = re.compile(r" {2,}")
    _CONTROL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
    _HYPHEN_BREAK = re.compile(r"-\n(\w)")

    @classmethod
    def normalize(cls, text: str) -> str:
        # Unicode normalization
        text = unicodedata.normalize("NFKC", text)
        # Remove control characters
        text = cls._CONTROL_CHARS.sub("", text)
        # Fix hyphenated line breaks (common in PDFs)
        text = cls._HYPHEN_BREAK.sub(r"\1", text)
        # Collapse excessive whitespace
        text = cls._MULTI_SPACE.sub(" ", text)
        # Collapse excessive newlines but keep paragraph breaks
        text = cls._MULTI_NEWLINE.sub("\n\n", text)
        return text.strip()


# ─── Base Parser ──────────────────────────────────────────────────────────────

class BaseParser(ABC):
    def __init__(self) -> None:
        self.normalizer = TextNormalizer()

    @abstractmethod
    async def parse(self, file_path: Path) -> tuple[str, Optional[int]]:
        """
        Parse a document file.
        Returns: (normalized_text, page_count)
        """

    def _clean(self, text: str) -> str:
        return self.normalizer.normalize(text)


# ─── PDF Parser ───────────────────────────────────────────────────────────────

class PDFParser(BaseParser):
    async def parse(self, file_path: Path) -> tuple[str, Optional[int]]:
        try:
            import pypdf
            import asyncio

            def _extract() -> tuple[str, int]:
                reader = pypdf.PdfReader(str(file_path))
                pages = []
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        pages.append(page_text)
                return "\n\n".join(pages), len(reader.pages)

            loop = asyncio.get_event_loop()
            text, page_count = await loop.run_in_executor(None, _extract)
            return self._clean(text), page_count
        except Exception as exc:
            raise IngestionError(f"PDF parsing failed: {exc}", {"file": str(file_path)}) from exc


# ─── DOCX Parser ─────────────────────────────────────────────────────────────

class DOCXParser(BaseParser):
    async def parse(self, file_path: Path) -> tuple[str, Optional[int]]:
        try:
            import asyncio
            from docx import Document

            def _extract() -> str:
                doc = Document(str(file_path))
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                # Also extract table content
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(
                            cell.text.strip() for cell in row.cells if cell.text.strip()
                        )
                        if row_text:
                            paragraphs.append(row_text)
                return "\n\n".join(paragraphs)

            loop = asyncio.get_event_loop()
            text = await loop.run_in_executor(None, _extract)
            return self._clean(text), None
        except Exception as exc:
            raise IngestionError(f"DOCX parsing failed: {exc}", {"file": str(file_path)}) from exc


# ─── TXT Parser ───────────────────────────────────────────────────────────────

class TXTParser(BaseParser):
    async def parse(self, file_path: Path) -> tuple[str, Optional[int]]:
        try:
            async with aiofiles.open(file_path, encoding="utf-8", errors="replace") as f:
                text = await f.read()
            return self._clean(text), None
        except Exception as exc:
            raise IngestionError(f"TXT parsing failed: {exc}", {"file": str(file_path)}) from exc


# ─── Markdown Parser ──────────────────────────────────────────────────────────

class MarkdownParser(BaseParser):
    _FRONTMATTER = re.compile(r"^---\n.*?\n---\n", re.DOTALL)
    _CODE_BLOCK = re.compile(r"```[\w]*\n(.*?)```", re.DOTALL)
    _INLINE_CODE = re.compile(r"`([^`]+)`")
    _MD_LINK = re.compile(r"\[([^\]]+)\]\([^\)]+\)")
    _MD_IMAGE = re.compile(r"!\[[^\]]*\]\([^\)]+\)")
    _MD_HEADING = re.compile(r"^#{1,6}\s+", re.MULTILINE)

    async def parse(self, file_path: Path) -> tuple[str, Optional[int]]:
        try:
            async with aiofiles.open(file_path, encoding="utf-8", errors="replace") as f:
                raw = await f.read()

            # Strip frontmatter
            text = self._FRONTMATTER.sub("", raw)
            # Remove image tags
            text = self._MD_IMAGE.sub("", text)
            # Preserve code blocks content
            text = self._CODE_BLOCK.sub(r"\1", text)
            # Simplify inline code
            text = self._INLINE_CODE.sub(r"\1", text)
            # Strip heading markers but keep text
            text = self._MD_HEADING.sub("", text)
            # Simplify links to just anchor text
            text = self._MD_LINK.sub(r"\1", text)

            return self._clean(text), None
        except Exception as exc:
            raise IngestionError(f"Markdown parsing failed: {exc}", {"file": str(file_path)}) from exc


# ─── Parser Factory ───────────────────────────────────────────────────────────

class ParserFactory:
    _PARSERS: dict[FileType, type[BaseParser]] = {
        FileType.PDF: PDFParser,
        FileType.DOCX: DOCXParser,
        FileType.TXT: TXTParser,
        FileType.MARKDOWN: MarkdownParser,
    }

    _EXTENSION_MAP: dict[str, FileType] = {
        ".pdf": FileType.PDF,
        ".docx": FileType.DOCX,
        ".doc": FileType.DOCX,
        ".txt": FileType.TXT,
        ".md": FileType.MARKDOWN,
        ".markdown": FileType.MARKDOWN,
    }

    @classmethod
    def get_file_type(cls, file_name: str) -> FileType:
        ext = Path(file_name).suffix.lower()
        if ext not in cls._EXTENSION_MAP:
            raise UnsupportedFileTypeError(
                f"File type '{ext}' is not supported. "
                f"Supported: {list(cls._EXTENSION_MAP.keys())}"
            )
        return cls._EXTENSION_MAP[ext]

    @classmethod
    def get_parser(cls, file_type: FileType) -> BaseParser:
        return cls._PARSERS[file_type]()
